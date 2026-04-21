"""
Document Loader Service
Handles extraction of text from PDF, DOCX, and TXT files with edge case handling.
"""

import os
from pathlib import Path
from typing import Optional
from pypdf import PdfReader
from pypdf.errors import PdfReadError
from docx import Document
from docx.opc.exceptions import PackageNotFoundError


class DocumentLoaderError(Exception):
    """Base exception for document loading errors."""
    pass


class EmptyFileError(DocumentLoaderError):
    """Raised when a file is empty or contains no extractable text."""
    pass


class CorruptedFileError(DocumentLoaderError):
    """Raised when a file is corrupted or cannot be read."""
    pass


class UnsupportedFileTypeError(DocumentLoaderError):
    """Raised when file type is not supported."""
    pass


class ScannedPDFError(DocumentLoaderError):
    """Raised when a PDF appears to be scanned (image-based) without OCR capability."""
    pass


class DocumentLoader:
    """
    Service for loading and extracting text from various document formats.
    
    Supported formats:
    - PDF (.pdf)
    - Word Documents (.docx)
    - Text Files (.txt)
    """
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt'}
    
    def __init__(self):
        """Initialize the DocumentLoader."""
        pass
    
    def load(self, file_path: str) -> str:
        """
        Load and extract text from a document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content as a string
            
        Raises:
            FileNotFoundError: If the file doesn't exist
            UnsupportedFileTypeError: If the file type is not supported
            EmptyFileError: If the file is empty or contains no text
            CorruptedFileError: If the file is corrupted
            ScannedPDFError: If PDF is scanned and OCR is not available
        """
        # Validate file exists
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file size
        if path.stat().st_size == 0:
            raise EmptyFileError(f"File is empty: {file_path}")
        
        # Get file extension
        extension = path.suffix.lower()
        
        # Validate file type
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(
                f"Unsupported file type: {extension}. "
                f"Supported types: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )
        
        # Route to appropriate loader
        if extension == '.pdf':
            return self._load_pdf(file_path)
        elif extension == '.docx':
            return self._load_docx(file_path)
        elif extension == '.txt':
            return self._load_txt(file_path)
    
    def _load_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
            
        Raises:
            CorruptedFileError: If PDF is corrupted
            EmptyFileError: If PDF contains no extractable text
            ScannedPDFError: If PDF appears to be scanned
        """
        try:
            reader = PdfReader(file_path)
            
            # Check if PDF has pages
            if len(reader.pages) == 0:
                raise EmptyFileError("PDF contains no pages")
            
            # Extract text from all pages
            text_content = []
            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
                except Exception as e:
                    # Continue with other pages if one fails
                    print(f"Warning: Could not extract text from page {page_num + 1}: {str(e)}")
                    continue
            
            # Join all extracted text
            full_text = "\n\n".join(text_content).strip()
            
            # Check if we got any text
            if not full_text:
                raise ScannedPDFError(
                    "PDF appears to be scanned or image-based. "
                    "No text could be extracted. OCR is not currently supported."
                )
            
            # Check if text is suspiciously short (might indicate scanned PDF)
            if len(full_text) < 50 and len(reader.pages) > 1:
                raise ScannedPDFError(
                    "PDF may be scanned or contain mostly images. "
                    "Very little text was extracted. OCR is not currently supported."
                )
            
            return full_text
            
        except PdfReadError as e:
            raise CorruptedFileError(f"PDF file is corrupted or invalid: {str(e)}")
        except (ScannedPDFError, EmptyFileError):
            # Re-raise our custom errors
            raise
        except Exception as e:
            raise CorruptedFileError(f"Error reading PDF file: {str(e)}")
    
    def _load_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
            
        Raises:
            CorruptedFileError: If DOCX is corrupted
            EmptyFileError: If DOCX contains no text
        """
        try:
            doc = Document(file_path)
            
            # Extract text from all paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            # Join all text
            full_text = "\n\n".join(text_content).strip()
            
            # Check if we got any text
            if not full_text:
                raise EmptyFileError("DOCX file contains no extractable text")
            
            return full_text
            
        except PackageNotFoundError as e:
            raise CorruptedFileError(f"DOCX file is corrupted or invalid: {str(e)}")
        except EmptyFileError:
            raise
        except Exception as e:
            raise CorruptedFileError(f"Error reading DOCX file: {str(e)}")
    
    def _load_txt(self, file_path: str) -> str:
        """
        Extract text from a TXT file.
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            File content as text
            
        Raises:
            EmptyFileError: If file contains no text
            CorruptedFileError: If file cannot be read
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read().strip()
                    
                    # If we successfully read the file, check if it has content
                    if not text:
                        raise EmptyFileError("TXT file contains no text")
                    
                    return text
                    
                except UnicodeDecodeError:
                    # Try next encoding
                    continue
            
            # If all encodings failed
            raise CorruptedFileError(
                "Could not decode text file. Tried encodings: " + ", ".join(encodings)
            )
            
        except EmptyFileError:
            raise
        except CorruptedFileError:
            raise
        except Exception as e:
            raise CorruptedFileError(f"Error reading TXT file: {str(e)}")


# Convenience function for quick usage
def load_document(file_path: str) -> str:
    """
    Convenience function to load a document.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Extracted text content
    """
    loader = DocumentLoader()
    return loader.load(file_path)
